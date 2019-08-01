import json
import pickle

import docx
import pandas as pd
from openpyxl import load_workbook


def save_pkl_file(save_data, file_name):
    """
    save data to pkl file
    :param save_data: array data for saving
    :param file_name: filename for saving, like '/path/filename.pkl'
    :return: None

    Other Method: Only DataFrame can use to_pickle() directly
        pd.DataFrame().to_pickle('/path/filename.pkl')
    """
    p = open(file_name, 'wb')
    pickle.dump(save_data, p)
    p.close()
    print('**************** Save Data To', file_name, 'Finished **************')


def load_pkl_file(file_name):
    """
    load data from pkl file
    :param file_name: pkl filename, like '/path/filename.pkl'
    :return: array data

    Other Method: Return the same type as data in pkl file
        data = pd.read_pickle('/path/filename.pkl')
    """
    p = open(file_name, 'rb')
    data = pickle.load(p, encoding='utf-8')
    p.close()
    print('*************** Load Data From', file_name, 'Finished *************')
    return data


def load_multi_pkl_file(file_name_list):
    """
    load data from multi pkl files
    :param file_name_list: filenames list, like ['/path/filename1.pkl', '/path/filename2.pkl', ...]
    :return: data list following the filename's order
    """
    all_data = []
    for file_name in file_name_list:
        p = open(file_name, 'rb')
        data = pickle.load(p, encoding='utf-8')
        all_data.append(data)
        p.close()
    print('*************** Load Data From Multi-pkl Files Finished *************')
    return all_data


def write_json_file(filename, data):
    """
    write data to json file
    :param filename: json filename, like '/path/filename.json'
    :param data: data for writing
    :return: None
    """
    fp = open(filename, mode='w', encoding='utf-8')
    json.dump(data, fp, ensure_ascii=False)
    fp.close()


def write_txt_file(filename, data_list):
    """
    write data to text file
    :param filename: txt filename, like '/path/filename.txt'
    :param data_list: data list for writing, like '[title, date, content, ...]'
    :return: None
    """
    fp = open(filename, mode='w', encoding='utf-8')
    for data in data_list:
        fp.write(data)
        fp.write('\n\n')
    fp.close()


def write_word_file(filename, title, data_list):
    """
    write text data to word file
    :param filename: word filename, like '/path/filename.doc' or '/path/filename.docx'
    :param title: title for word file, or maybe None
    :param data_list: paragraphs of word content
    :return: None
    """
    doc = docx.Document()
    if title:
        doc.add_heading(title, 0)
    for data in data_list:
        doc.add_paragraph(data)
    doc.save(filename)


def write_excel_file(filename, data_list, sheet_name='Sheet1', columns=None):
    """
    write list data to excel file, supporting add new sheet
    :param filename: excel filename, like '/path/filename.xls'
    :param data_list: list data for saving
    :param sheet_name: excel sheet name, default 'Sheet1'
    :param columns: excel column names
    :return: None
    """
    writer = pd.ExcelWriter(filename)
    frame = pd.DataFrame(data_list, columns=columns)
    book = load_workbook(writer.path)
    writer.book = book
    frame.to_excel(excel_writer=writer, sheet_name=sheet_name, index=None)
    writer.close()
